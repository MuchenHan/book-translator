#!/usr/bin/env python3
"""
Build a formatted docx from page-marked text (=== p.XX === markers).
Usage: python3 build_docx.py <input.txt> <output.docx> [options]

Options:
  --title       Book/chapter title
  --subtitle    Subtitle or original title
  --author      Author name
  --year        Publication year
  --font        Main font (default: Yu Mincho)
  --font-size   Body font size in pt (default: 11)
  --terms-file  Path to terminology TSV/CSV (english<TAB>japanese per line)
  --sections    Comma-separated section titles to detect as headings
"""
import re, sys, argparse, csv
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("目次")
    run.bold = True
    run.font.size = Pt(14)

    p2 = doc.add_paragraph()
    run2 = p2.add_run()
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    run2._r.append(fc1)

    run3 = p2.add_run()
    it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = ' TOC \\o "1-3" \\h \\z \\u '
    run3._r.append(it)

    run4 = p2.add_run()
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'separate')
    run4._r.append(fc2)

    run5 = p2.add_run("（右クリック → フィールドの更新 で目次が生成されます）")
    run5.font.color.rgb = RGBColor(128, 128, 128)
    run5.font.size = Pt(9)

    run6 = p2.add_run()
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'end')
    run6._r.append(fc3)

    doc.add_page_break()


def load_terms(terms_file):
    terms = []
    with open(terms_file, 'r') as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            # Support TSV, CSV, or "eng → jpn" format
            if '\t' in line:
                parts = line.split('\t', 1)
            elif '→' in line:
                parts = [p.strip() for p in line.split('→', 1)]
            elif ',' in line:
                parts = line.split(',', 1)
            else:
                continue
            if len(parts) == 2:
                # Strip markdown list markers
                eng = parts[0].lstrip('- ').strip()
                jpn = parts[1].strip()
                terms.append((eng, jpn))
    return terms


def build_docx(input_txt, output_docx, title="", subtitle="", author="", year="",
               font_name="Yu Mincho", font_size=11, terms_file=None, sections=None):
    doc = Document()

    # Default style
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title page
    if title:
        for _ in range(6):
            doc.add_paragraph()
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = tp.add_run(title)
        r.bold = True
        r.font.size = Pt(24)

        if subtitle:
            sp = doc.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = sp.add_run(subtitle)
            r.italic = True
            r.font.size = Pt(16)
            r.font.name = 'Times New Roman'

        if author:
            doc.add_paragraph()
            ap = doc.add_paragraph()
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = ap.add_run(author)
            r.font.size = Pt(14)

        if year:
            yp = doc.add_paragraph()
            yp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = yp.add_run(year)
            r.font.size = Pt(12)

        doc.add_page_break()

    # TOC
    add_toc(doc)

    # Build section header set for heading detection
    section_set = {}
    if sections:
        for s in sections:
            s = s.strip()
            if ':' in s:
                name, level = s.rsplit(':', 1)
                section_set[name.strip()] = int(level)
            else:
                section_set[s] = 2

    # Read and split by page markers
    with open(input_txt, 'r') as f:
        text = f.read()

    pages = re.split(r'=== p\.(\d+) ===', text)
    last_page = 0

    for i in range(1, len(pages), 2):
        page_num = int(pages[i])
        content = pages[i + 1].strip()
        last_page = page_num

        # Page header
        ph = doc.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = ph.add_run(f'— p.{page_num} —')
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(128, 128, 128)
        r.font.name = 'Times New Roman'

        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue

            # Check section headers
            is_header = False
            for header, level in section_set.items():
                if line == header:
                    h = doc.add_heading(line, level=level)
                    for run in h.runs:
                        run.font.name = font_name
                    is_header = True
                    break

            if not is_header:
                p = doc.add_paragraph(line)
                p.paragraph_format.first_line_indent = Pt(22)

        doc.add_page_break()

    # Terminology appendix
    if terms_file:
        terms = load_terms(terms_file)
        if terms:
            h = doc.add_heading('術語対照表', level=1)
            for run in h.runs:
                run.font.name = font_name

            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'English'
            hdr[1].text = '日本語'
            for cell in hdr:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(10)

            for eng, jpn in terms:
                row = table.add_row().cells
                row[0].text = eng
                row[1].text = jpn
                for cell in row:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(10)

    doc.save(output_docx)
    print(f"Saved: {output_docx} ({last_page - int(pages[1]) + 1} pages)")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Build formatted docx from page-marked text")
    parser.add_argument("input", help="Input text file with === p.XX === markers")
    parser.add_argument("output", help="Output docx file path")
    parser.add_argument("--title", default="")
    parser.add_argument("--subtitle", default="")
    parser.add_argument("--author", default="")
    parser.add_argument("--year", default="")
    parser.add_argument("--font", default="Yu Mincho")
    parser.add_argument("--font-size", type=int, default=11)
    parser.add_argument("--terms-file", default=None, help="Terminology file (TSV/CSV/arrow format)")
    parser.add_argument("--sections", default=None, help="Section titles, comma-separated. Append :N for heading level (default 2)")
    args = parser.parse_args()

    sections = args.sections.split(',') if args.sections else None
    build_docx(args.input, args.output, args.title, args.subtitle, args.author, args.year,
               args.font, args.font_size, args.terms_file, sections)
